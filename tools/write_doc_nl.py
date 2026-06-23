from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

PURPLE = RGBColor(0x6B, 0x21, 0xA8)
GRAY   = RGBColor(0x6B, 0x72, 0x80)

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = PURPLE
    return p

def h2(text):
    return doc.add_heading(text, level=2)

def h3(text):
    return doc.add_heading(text, level=3)

def body(text):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(10.5)
    return p

def bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.runs[0].font.size = Pt(10.5)
    return p

def note(text):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = GRAY
    return p

def gap():
    doc.add_paragraph()

# -----------------------------------------------------------------------
# TITLE
# -----------------------------------------------------------------------
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("LOFI Artist Intelligence Platform")
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = PURPLE

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
s.add_run("Technische samenvatting  |  Juni 2026  |  Intern gebruik").font.size = Pt(9)

gap()

# -----------------------------------------------------------------------
# 1. DOEL
# -----------------------------------------------------------------------
h1("1. Wat doet het systeem?")
body(
    "Het LOFI Artist Intelligence Platform is een intern data- en dashboardsysteem dat twee "
    "boekingsvragen beantwoordt:"
)
bullet("Wie moeten we nu al in de gaten houden? -- opkomende artiesten herkennen 6-18 maanden voor de rest van de markt.")
bullet("Kan deze artiest tickets verkopen in Amsterdam? -- data-onderbouwde beoordeling voor boekingsbeslissingen.")
body(
    "Kernprincipe: groeiversnelling telt zwaarder dan absolute grootte. "
    "De tweede afgeleide van Spotify-luisteraarsgroei (versnelling) is het primaire signaal "
    "door het hele systeem: in scores, anomaliedetectie en het ML-model."
)
gap()

# -----------------------------------------------------------------------
# 2. PIPELINE OVERZICHT
# -----------------------------------------------------------------------
h1("2. Pipeline -- hoe stroomt data door het systeem?")
body(
    "De pipeline loopt van links naar rechts: Scrapers verzamelen ruwe data uit externe bronnen "
    "en slaan die op in Supabase. De Scoring Engine berekent nachtelijk vijf scores per artiest "
    "op basis van die data. Het XGBoost-model maakt groeiverwachtingen op basis van historische "
    "tijdreeksen. Het dashboard leest alles uit Supabase en toont het aan de gebruiker. "
    "Validatie-events (eerste Ibiza, eerste Boiler Room etc.) worden automatisch gedetecteerd "
    "door de live scrapers en opgeslagen als mijlpalen."
)
gap()

# -----------------------------------------------------------------------
# 3. DATABRONNEN EN SCRAPERS
# -----------------------------------------------------------------------
h1("3. Databronnen en Scrapers")

# 3.1 Chartmetric
h2("3.1 Chartmetric API -- kern van de artiestdata")
body(
    "Chartmetric is de primaire databron. Via hun REST API worden voor elke artiest maandelijkse "
    "snapshots opgehaald van Spotify-luisteraars en -volgers, Instagram-volgers, TikTok-volgers "
    "en -likes, SoundCloud-volgers, YouTube-kanaalstatistieken, Deezer-fans en Last.fm-luisteraars. "
    "Daarnaast worden de CPP-score (Chartmetric's eigen populariteitsindex op basis van "
    "cross-platform groei), editorial playlist-plaatsingen en top tracks opgehaald."
)
bullet("Opslag: ruwe tijdreeksdata in artist_chartmetric.cm_timeseries als JSONB. Platte huidige "
       "waarden in artist_chartmetric_flat voor snelle dashboard-queries.")
bullet("Backfill: bij het toevoegen van een nieuwe artiest haalt backfill_booked.py de volledige "
       "historische tijdreeks op, zodat het ML-model direct over historische data beschikt.")
bullet("ML features: na elke backfill berekent het systeem automatisch ml_features -- "
       "30d/90d/180d groeipercentages en groeiversnelling -- en slaat die op naast de tijdreeks.")

# 3.2 Resident Advisor
h2("3.2 Resident Advisor -- eventgeschiedenis en scenevalidatie")
body(
    "De RA-scraper haalt via de RA GraphQL API eventhistorie op per artiest: datum, venue, stad, "
    "land, capaciteit, lineup-grootte en of de artiest headliner was. De data wordt opgeslagen "
    "in de artist_ra.events JSONB-kolom (ruwe scrape) en in de genormaliseerde ra_events-tabel "
    "(nachtelijkse run van scrape_ra_events.py)."
)
bullet("Gebruik in scores: RA-events zijn input voor het Scene Signaal -- Ibiza-optredens (Amnesia, Hi, DC-10), "
       "Awakenings, Thuishaven en andere validatielocaties verhogen de scene-score.")
bullet("Gebruik in dashboard: RA-eventlijst in het artiestprofiel, filterbaar op NL / Ibiza / headliner.")
bullet("Mijlpalen: eerste Ibiza-optreden, eerste headlineset in een zaal van 500+ worden automatisch "
       "opgeslagen als validatie-events.")

# 3.3 Partyflock
h2("3.3 Partyflock -- Amsterdam/NL-vraag")
body(
    "Partyflock is het primaire signaal voor Amsterdam-vraag. De scraper haalt op: aantal fans "
    "(pf_fans), totaal aantal optredens, aankomende en afgelopen optredens, en een JSONB-array "
    "van alle events met datum, venue en stad. 912 van de 912 artiesten hebben een Partyflock-rij."
)
bullet("Gebruik in NL-score: het aantal NL-events (filtering op stad/land) en de pf_fans-teller "
       "zijn directe input voor de NL / Amsterdam-score op het artiestprofiel.")
bullet("Gebruik in dashboard: Partyflock-tab in Shows toont alle NL-events, met uitklapbare volledige lijst.")

# 3.4 Last.fm
h2("3.4 Last.fm -- genre-tags en vergelijkbare artiesten")
body(
    "De Last.fm-scraper haalt genre-tags, luisteraantallen, afspeeltelling en vergelijkbare "
    "artiesten op. De tags worden opgeslagen als JSONB in artist_lastfm.tags."
)
bullet("Gebruik in Genre Trends pagina: tags worden geexplodeerd en geaggregeerd per tag "
       "om te laten zien welke genres groeien in artiestenaantal en gemiddeld luisteraarsvolume.")
bullet("Gebruik in LOFI Feel score: de LOFI Feel-taxonomie (scoring/lofi_feel_taxonomy.yaml) "
       "koppelt genre-tags aan LOFI programmapijlers (dark techno, driving, melodic, organic, minimal). "
       "Hoe meer tags matchen, hoe hoger de LOFI Fit-score.")
bullet("Vergelijkbare artiesten: lfm_similar_artists wordt getoond in het artiestprofiel.")

# 3.5 YouTube Monitor
h2("3.5 YouTube Monitor -- live detectie van sets (elke 30 min)")
body(
    "De YouTube Monitor draait elke 30 minuten via GitHub Actions. Per kanaal haalt hij de "
    "laatste 10 videos op via de YouTube Data API v3 (playlistItems endpoint). "
    "Kanalen: Boiler Room, HOR Berlin, Mixmag, The Lot Radio, Book Club Radio, Rinse FM, BE-AT.TV."
)
bullet("Artiestherkenning: per kanaal is een regex gedefinieerd in channel_config.py die de "
       "artiestnaam uit de videotitel extraheert. B2B-sets worden gesplitst op ' b2b ', ' & ', ' x '.")
bullet("Matching: gevonden namen worden fuzzy-gematcht (rapidfuzz, token_set_ratio >= 85) "
       "tegen alle artiesten in de database. Onbekende namen gaan naar discovery_queue.")
bullet("Velocity: views worden elke 30 min opgeslagen in youtube_snapshots. "
       "Velocity = (nieuwe views - oude views) / verstreken uren. "
       "Trending-drempel verschilt per kanaal (Boiler Room: 8.000 v/u, HOR: 3.000 v/u).")
bullet("Mijlpalen: eerste optreden op een kanaal wordt automatisch opgeslagen "
       "(first_boiler_room, first_hor_berlin etc.) in validation_events.")
bullet("Dashboard: YouTube Sets-pagina toont het thumbnail-raster met views, kanaalkleur en trending-badge.")

# 3.6 Beatport
h2("3.6 Beatport Charts -- commerciele chartposities (elke 6 uur)")
body(
    "De Beatport-scraper haalt de top 100 HTML-pagina op voor vijf genres: "
    "Techno (Peak Time), Tech House, Melodic House & Techno, Afro House, Organic House. "
    "Trackdata wordt geextraheerd uit het __NEXT_DATA__ JSON-blob in de paginabron "
    "en opgeslagen in beatport_chart_entries."
)
bullet("Artiestkoppeling: artiesten in de chart worden gematcht tegen de database. "
       "Onbekende artiesten worden ook opgeslagen (artist_id = NULL) voor trendanalyse.")
bullet("Mijlpalen: eerste Beatport-notering, eerste top 10, eerste nummer 1 worden "
       "automatisch opgeslagen als validatie-events.")

# 3.7 RA Podcast
h2("3.7 RA Podcast -- branche-validatie (dagelijks)")
body(
    "De RA Podcast-scraper gebruikt de RA GraphQL API (ra.co/graphql) om dagelijks de "
    "meest recente podcast-afleveringen op te halen. Let op: de RA GraphQL API geeft "
    "maximaal 10 resultaten terug -- hogere limieten leveren een lege array op."
)
bullet("Mijlpalen: eerste RA Podcast-aflevering wordt opgeslagen als first_ra_podcast "
       "in validation_events, inclusief afleveringstitel en URL in de details-JSONB.")

# 3.8 BBC Essential Mix
h2("3.8 BBC Essential Mix -- prestige-validatie (wekelijks)")
body(
    "De BBC-scraper gebruikt de BBC Sounds RMS API om wekelijks te controleren op "
    "nieuwe Essential Mix-afleveringen. De Essential Mix is een van de meest prestigieuze "
    "radio-optredens in de elektronische muziek."
)
bullet("Mijlpalen: eerste Essential Mix wordt opgeslagen als first_bbc_essential_mix.")
gap()

# -----------------------------------------------------------------------
# 4. SCORING ENGINE
# -----------------------------------------------------------------------
h1("4. Scoring Engine -- vijf scores per artiest")
body(
    "Elke nacht herberekent scoring/lofi_scorer.py vijf scores (0-100) voor alle artiesten. "
    "De scores worden opgeslagen in artist_chartmetric_scores met een tijdstempel, "
    "zodat de scoregeschiedenis bewaard blijft voor ML-training."
)

h2("Score 1: Momentum (cross-platform buzz nu)")
body("Momentum = 35% Spotify 30d groei + 30% cross-platform momentum + 20% TikTok 30d + 15% Instagram 30d")
body(
    "Alle groeipercentages gaan door een logistische (tanh) mapping: 0% groei -> score 50, "
    "+30% groei -> score ~75, -30% groei -> score ~25. Dit voorkomt dat uitschieters "
    "de score domineren."
)

h2("Score 2: Groei (versnelling als primair signaal)")
body("Groei = 50% groeiversnelling + 30% Spotify 30d + 20% CPP 30d trend")
body(
    "Groeiversnelling (sp_listeners_accel) is de tweede afgeleide: hoeveel sneller of langzamer "
    "groeit de artiest nu ten opzichte van de voorgaande periode? "
    "Formule: accel_score = sp_listeners_accel * 2 + 50 (gecentreerd op 50). "
    "Dit is het kernprincipe van het platform: een artiest die versnelt is interessanter dan "
    "een grote artiest die stabiel blijft."
)

h2("Score 3: Marktpositie (huidige omvang)")
body("Marktpositie = 40% CM-rank + 30% CPP-score + 20% Spotify-luisteraars + 10% fanbase-rank")
body("Hoe lager de CM-artiestenrang (ranking onder alle artiesten worldwide), hoe hoger de score.")

h2("Score 4: Potentieel (waar gaat dit naartoe?)")
body("Potentieel = 40% groeiversnelling + 30% CPP 90d trend + 30% Spotify 180d trend")
body(
    "Het Potentieel-score kijkt verder vooruit dan Momentum: 180 dagen groeitrend "
    "en de langetermijn CPP-beweging bepalen of een artiest een structurele opwaartse "
    "lijn laat zien."
)

h2("Score 5: Data (betrouwbaarheid)")
body(
    "Elke ingevulde databron voegt ~16-20 punten toe: Spotify, Instagram, TikTok, "
    "SoundCloud, YouTube, Last.fm. Een artiest met data op 5 van de 6 platforms scoort ~83. "
    "Een artiest met alleen Spotify-data scoort ~20. Deze score is zichtbaar als "
    "'Confidence' en bepaalt mede hoeveel gewicht de andere scores krijgen."
)

h2("Booking Signalen (composiet)")
body(
    "Naast de vijf individuele scores berekent het dashboard drie composiet booking-signalen:"
)
bullet("GROEI SIGNAAL (40%) -- XGBoost ML-voorspelling van CPP-groei in de komende 90 dagen, "
       "omgezet naar 0-100 (0% groei = 50, +50% groei = 100, -50% groei = 0).")
bullet("SCENE SIGNAAL (35%) -- combinatie van validatiescore (Ibiza, Awakenings, Beatport, "
       "RA Podcast, BBC etc.), NL-score (Partyflock fans + NL-events) en RA-eventgeschiedenis.")
bullet("LOFI FIT (25%) -- genrematch met de LOFI Feel-taxonomie (YAML-config). "
       "Welke LOFI programmapijlers (dark techno, driving, melodic) matchen met de genre-tags?")
gap()

# -----------------------------------------------------------------------
# 5. XGBOOST MODEL
# -----------------------------------------------------------------------
h1("5. XGBoost Groeivoorspelling")
body(
    "Het ML-model voorspelt de procentuele verandering in Chartmetric CPP-score "
    "(audience_index) over de komende 90 dagen. Dit is een echte forward-looking "
    "voorspelling: de target-waarde wordt berekend als shift(-90) op de historische "
    "tijdreeks -- de CPP-waarde die 90 dagen later daadwerkelijk is geobserveerd."
)

h2("Training (ml/train_growth_model.py)")
bullet("Features: voor elk historisch tijdstip per artiest worden 7d/30d/90d-groeisnelheden, "
       "groeiversnelling, voortschrijdend gemiddelde en standaardafwijking, "
       "variatiecoefficient en seizoenspatroon berekend.")
bullet("Target: 90 dagen voorwaartse CPP-groei in procenten. Toekomstige waarden die nog "
       "niet bekend zijn worden uitgefilterd -- er is geen datalekkage.")
bullet("Cross-validatie: GroupShuffleSplit op artist_id. Artiest A zit nooit tegelijk "
       "in de trainset en de validatieset. Dit voorkomt dat het model leert van de "
       "specifieke artistenpatronen in plaats van generaliseerbare groeipatronen.")
bullet("Model: XGBRegressor. Opgeslagen als ml/models/growth_predictor.json. "
       "Metrische output: MAE, R2, Spearman-correlatie in model_meta.json.")

h2("Voorspelling (ml/bulk_predict_to_supabase.py)")
body(
    "Draait nachtelijks. Laadt growth_predictor.json, berekent features voor alle 912 "
    "artiesten op basis van de meest recente cm_timeseries-snapshot, en schrijft de "
    "voorspelling naar xgboost_predictions. "
    "Per artiest wordt ook missing_pct bijgehouden: hoeveel van de features ontbreken "
    "vanwege ontbrekende tijdreeksdata. Bij >40% ontbrekende features toont het dashboard "
    "een lage-betrouwbaarheidswaarschuwing."
)

h2("Gebruik in het dashboard")
bullet("Groei Leaderboard: standaard gesorteerd op predicted_growth_90d, aflopend.")
bullet("Artiestkaartjes: XGBoost-groeipercentage rechtsonder op elke kaart.")
bullet("Artiestprofiel: GROEI SIGNAAL composiet-score op basis van de voorspelling.")
gap()

# -----------------------------------------------------------------------
# 6. SIMILARITY / ARTIEST RECOMMENDER
# -----------------------------------------------------------------------
h1("6. Similarity -- Artiest Recommender")
body(
    "De Artiest Recommender combineert twee signalen om vergelijkbare artiesten te vinden:"
)

h2("Signaal 1: FAISS vector similarity")
body(
    "Van elke artiest wordt een feature-vector gebouwd op basis van platform-aantallen "
    "(Spotify-luisteraars, Instagram-volgers etc.), groeipercentages en genre-embeddings "
    "(one-hot encoding van Last.fm-tags). Deze vectoren worden opgeslagen in "
    "data/faiss.index via Facebook FAISS (Flat L2-index)."
)
body(
    "Bij een query wordt de vector van de input-artiest opgezocht en worden de K "
    "dichtstbijzijnde vectoren in de index teruggegeven -- artiesten met een vergelijkbaar "
    "geluid, vergelijkbare omvang en vergelijkbare groeisnelheid."
)

h2("Signaal 2: RA co-lineup data")
body(
    "De lineup_recommender parst alle RA-event lineups en bouwt een co-voorkomstmatrix: "
    "welke artiesten staan het vaakst samen op dezelfde poster? "
    "Co-voorkomsten worden gewogen naar venuetier: Awakenings weegt zwaarder dan een "
    "kleine clubnacht."
)
body(
    "Bij een query worden de top-N artiesten teruggegeven die het vaakst samen "
    "met de input-artiest in RA-lineups voorkomen, gesorteerd op gewogen co-voorkomstfrequentie."
)

h2("Combinatie in het dashboard")
body(
    "De Artiest Recommender-pagina toont beide lijsten naast elkaar. "
    "In het artiestprofiel worden Last.fm Similar Artists en Chartmetric Related Artists "
    "als tekstlinks getoond onder 'Vergelijkbare artiesten'."
)
gap()

# -----------------------------------------------------------------------
# 7. DATABASE
# -----------------------------------------------------------------------
h1("7. Database -- Supabase, tinder-schema")
body(
    "Alle data staat in het tinder-schema van het Supabase-project. "
    "De koppelsleutel door het hele systeem is artist_id (UUID). "
    "Er zijn geen orphaned rows -- elke tabel is volledig gekoppeld."
)

h2("Kerntabellen")
bullet("artist_chartmetric_flat -- platte huidige waarden voor alle platforms. Primaire tabel voor dashboardqueries.")
bullet("artist_chartmetric -- tijdreeksen (cm_timeseries JSONB) en ML-features. Zware JSONB-payload, ~745/912 gevuld.")
bullet("artist_cm_extended -- rijke Chartmetric-data: Shazam, demographics, playlists, tracks, nieuws. 282/912 gevuld.")
bullet("ra_events -- genormaliseerde RA-eventhistorie per artiest. 133 artiesten in de genormaliseerde tabel, "
       "620 artiesten in de artist_ra.events JSONB-fallback.")
bullet("artist_partyflock -- Partyflock NL-data inclusief events-JSONB. 912/912 gevuld.")
bullet("artist_lastfm -- Last.fm tags, luisteraantallen, afspeeltelling. 912/912 gevuld.")
bullet("validation_events -- gevalideerde carriere-events (first_ibiza, first_headline_500, "
       "first_boiler_room, first_ra_podcast etc.) met datum, bron en bevestigingsstatus.")
bullet("xgboost_predictions -- ML-modeloutput: predicted_growth_90d, missing_pct, model_version, predicted_at.")
bullet("artist_chartmetric_scores -- scoregeschiedenis: momentum, groei, marktpositie, potentieel, "
       "confidence per artiest per dag.")
bullet("youtube_sets -- gescrapete YouTube-videos met thumbnail_url, view_count, view_velocity, is_trending.")
bullet("beatport_chart_entries -- Beatport top 100-posities per genre per artiest per datum.")
bullet("discovery_queue -- onbekende artiesten gevonden in trending content, wachtend op review.")
gap()

# -----------------------------------------------------------------------
# 8. AUTOMATISERING
# -----------------------------------------------------------------------
h1("8. Automatisering -- GitHub Actions")
bullet("Elke 30 min -- [LD] YouTube Monitor: trending sets detecteren op 7 kanalen.")
bullet("Elke 6 uur -- [LD] Beatport Charts: top 100 in 5 genres bijwerken.")
bullet("Dagelijks 06:00 -- RA Podcast checken.")
bullet("Dagelijks 02:00 -- Score-engine: vijf scores herberekenen voor alle 912 artiesten.")
bullet("Dagelijks (nacht) -- XGBoost bulk predict: voorspellingen bijwerken.")
bullet("Wekelijks (zondag) -- [LD] BBC Essential Mix: nieuwe uitzendingen detecteren.")
bullet("Op aanvraag -- Chartmetric backfill, XGBoost model retraining, nieuwe artiesten toevoegen.")
note("Vereiste GitHub Secrets: SUPABASE_URL, SUPABASE_KEY, YOUTUBE_API_KEY, CM_REFRESH_TOKEN.")
gap()

# -----------------------------------------------------------------------
# 9. BEKENDE GAPS
# -----------------------------------------------------------------------
h1("9. Bekende datahiaten")
bullet("NL publiek % -- Instagram/TikTok country breakdown beschikbaar voor 282/912 artiesten. Dashboard toont melding als het ontbreekt.")
bullet("Agency-vertegenwoordiging -- geen publieke API. Handmatige invoer gepland. Hoge prioriteit: agency-tier is een sterke leading indicator.")
bullet("Fan cities (Chartmetric) -- kolom bestaat maar is leeg voor alle artiesten. Nog niet aangeleverd door Chartmetric.")
bullet("F2F TV YouTube -- channel ID ontbreekt, wordt geskipt door de YouTube Monitor.")
bullet("LOFI intern kaartverkoop -- nog niet gekoppeld. Gepland voor latere fase als aanvullend signaal voor 'kan deze artiest tickets verkopen in Amsterdam?'.")
gap()

# FOOTER
f = doc.add_paragraph()
f.alignment = WD_ALIGN_PARAGRAPH.CENTER
f.add_run("LOFI Artist Intelligence Platform  |  LarsGroep/LOFI  |  Juni 2026").font.size = Pt(8)

doc.save(r"C:\Users\larsv\Desktop\LOFI repo NL samenvatting.docx")
print("Saved.")
